
from typing import List
from uuid import UUID
from scrc.enums.section import Section
from enum import Enum
from pathlib import Path
from sqlite3 import Connection
import sys
from scrc.preprocessors.abstract_preprocessor import AbstractPreprocessor
from scrc.utils.log_utils import get_logger
from scrc.utils.main_utils import get_config
from scrc.enums.section import Section
from docx import Document


class Color(Enum):
    HEADER = 7
    FACTS = 8
    CONSIDERATIONS = 10
    RULINGS = 16
    FOOTER = 3


class CoverageVerification(AbstractPreprocessor):
    """Ouputs docx files of court rulings with highlighted paragraphs of each section recognized aswell as the ruling outcome to data/verification folder. 
    To run simply remove your desired spider from coverage_verification.txt and run:
    
    python -m scrc.analyses.coverage_verification | python -m scrc.analyses.coverage_verification 1 (for rulings only)
    
    If you wish to increase the outputted document size, change the range in line 56 to a higher number."""
    
    def __init__(self, config: dict):
        super().__init__(config)
        self.logger = get_logger(__name__)
        self.logger_info = {
            'start': 'Started coverage verification',
            'finished': 'Finished coverage verification',
            'start_spider': 'Started coverage verification for spider',
            'finish_spider': 'Finished coverage verification for spider',
        }
        self.processed_file_path = self.progress_dir / "coverage_verification.txt"

    
    def start(self):
        """Starts the coverage verification. """
        spider_list, message = self.compute_remaining_spiders(self.processed_file_path)
        self.logger.info(message)
        engine = self.get_engine(self.db_scrc)
        with engine.connect() as conn:
            for spider in spider_list:
                self.write_document(spider, conn )
                self.logger.info(self.logger_info['finish_spider'] + ' ' + spider)
                self.mark_as_processed(self.processed_file_path, spider)
            self.logger.info(self.logger_info['finished'])
        
                
    def write_document(self, spider: str, conn: Connection):
        """Writes the document to the data/verification folder."""
        document = Document()
        decision_list = []
        for x in range(0, 50):  
            result = self.get_valid_decision(conn, spider, decision_list)
            self.write_to_doc(document, result, x)
        document.save(self.get_path(spider))
 
        
    def write_to_doc(self, document: Document, result: dict, x: int):
        """Writes the sections of the decision to the document. If the script is run with an argument, only the ruling outcome is written."""
        sections = result['wrapper']['sections']
        document.add_paragraph(f"{str(result['wrapper']['id'])}: {x + 1}")
        if len(sys.argv) > 1:
            p = document.add_paragraph()
            p.add_run(sections['RULINGS'])
            p = document.add_paragraph()
            r = p.add_run(result['ruling']).bold = True 
        else:
            sorted_list = sorted(list(sections.items()), key=lambda x: Section[x[0]].value)
            for element in (sorted_list):
                p = document.add_paragraph()
                r = p.add_run(f"{element[0]}: {element[1]}").font
                r.highlight_color = Color[element[0]].value
                if element[0] == 'RULINGS':
                    p = document.add_paragraph()
                    r = p.add_run(result['ruling']).bold = True 
         
    def get_valid_decision(self, conn: Connection, spider: str, decision_list: List[UUID]):
        """Returns a valid decision that has not been used before and has a ruling outcome."""
        MAX_LOOPS = 2000
        index = 0
        invalid_decision = True
        while invalid_decision:
            index += 1
            if index > MAX_LOOPS:
                self.logger.info("Did not find a valid decision after 10'000 tries, aborting")
                break
            result = conn.execute(self.random_decision_query(spider)).fetchone()
            if result[0] not in decision_list:
                rulings = self.get_ruling_outcome(str(result[0]), conn)
                if rulings:
                    invalid_decision = False
                    decision_list.append(result[0])
                    sections = self.get_sections(result, conn)
                    ruling_text = ""
                    for ruling in rulings:
                        ruling_text += ruling.text + " "
                    return {"wrapper": sections, "ruling": ruling_text}
          
            
    def get_sections(self, result: str, conn: Connection,):
        """Returns the sections of the decision."""
        section_dict = {'sections': {}, 'id': result[0]}
        result = conn.execute(self.section_query(result[0]))
        for res in result:
            section_dict['sections'][Section(res.section_type_id).name] = res.section_text
        return section_dict
    
    
    def get_ruling_outcome(self, decision_id: str, conn: Connection):
        return conn.execute(self.ruling_outcome_query(decision_id)).fetchall()
        
    
    def ruling_outcome_query(self, decision_id: str):
        """Returns the ruling outcome of a decision."""
        return (f"SELECT * FROM judgment "
                f"INNER JOIN judgment_map ON judgment.judgment_id = judgment_map.judgment_id "
                f"INNER JOIN decision ON judgment_map.decision_id = decision.decision_id "
                f"WHERE decision.decision_id = '{decision_id}'")

    def section_query(self, decision_id: str):
        """Returns all sections of a decision except for the full text and topic sections."""
        return (f"SELECT * FROM section "
                f"WHERE section.decision_id = '{decision_id}' "
                f"AND section.section_type_id != {Section.FULL_TEXT.value} "
                f"AND section.section_type_id != {Section.TOPIC.value} ")
        
    def random_decision_query(self, spider: str):
        """Returns a random decision for a given spider from the database."""
        return (f"SELECT * FROM decision "
                f"INNER JOIN chamber ON decision.chamber_id = chamber.chamber_id "
                f"INNER JOIN spider ON chamber.spider_id = spider.spider_id "
                f"WHERE spider.name = '{spider}' "
                f"ORDER BY RANDOM() LIMIT 1")
        
    def get_path(self, spider: str):
        filepath = Path(
        f'data/verification/{spider}.docx')
        filepath.parent.mkdir(parents=True, exist_ok=True)
        return filepath
        
        
if __name__ == '__main__':
    config = get_config()

    coverage_verification = CoverageVerification(config)
    coverage_verification.start()